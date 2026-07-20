from __future__ import annotations

import argparse
import os
from pathlib import Path
from typing import Final

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from korean_auto_translation import (
    DEFAULT_AUTO_TRANSLATIONS_PATH,
    DEFAULT_TRANSLATION_MODEL,
    load_auto_translations,
    save_auto_translations,
    translate_prose,
)


TRANSLATIONS: dict[str, str] = {
    "I study how evaluations of originality, authority, and credibility are formed, whose work those judgments elevate, and why inequality is greatest where standards are least clear.":
        "저는 독창성, 권위, 신뢰성에 대한 평가가 어떻게 형성되는지, 그 판단이 누구의 작업에 권위와 인정을 부여하는지, 그리고 기준이 가장 불명확한 곳에서 왜 불평등이 가장 크게 나타나는지 연구합니다.",
    "I study how evaluation determines whose work is recognized as original, authoritative, and credible, and why inequality concentrates most where evaluation is least standardized.":
        "평가 체계가 누구의 연구를 독창적이고 권위 있으며 신뢰할 만한 것으로 인정하는지, 그리고 평가 기준이 가장 덜 표준화된 영역에서 왜 불평등이 집중되는지 연구합니다.",
    "Jina Lee is a sociologist of evaluation, knowledge, and inequality. Her work asks: how do systems of evaluation determine whose contributions are recognized as original, authoritative, and credible? And why does inequality concentrate in the judgments least governed by clear standards? She pursues these questions across three arenas: the sciences; cultural canonization; and the AI systems now absorbing parts of scholarly judgment itself. Using computational text analysis, bibliometric analysis, and experiments, she finds that women and lower-status scholars are less often granted authority and recognition, and that this gap widens where standards are most ambiguous. Her research appears in the American Sociological Review, Gender & Society (forthcoming), and Poetics, among other venues.":
        "평가, 지식, 불평등을 연구하는 사회학자 이진아입니다. 평가 체계는 누구의 기여를 독창적이고 권위 있으며 신뢰할 만한 것으로 인정하는지 어떻게 결정할까요? 그리고 명확한 기준의 지배를 가장 덜 받는 판단에서 왜 불평등이 집중될까요? 이 질문을 과학, 문화적 정전화, 그리고 학술적 판단의 일부를 흡수하고 있는 AI 시스템이라는 세 영역에서 탐구합니다. 전산 텍스트 분석, 서지 분석, 실험을 활용해 여성 연구자와 지위가 낮은 연구자들이 권위와 인정을 덜 부여받으며, 그 격차가 기준이 가장 모호한 곳에서 확대된다는 점을 보여줍니다. 연구는 American Sociological Review, Gender & Society (게재 예정), Poetics를 비롯한 여러 학술지에 실렸습니다.",
    "Jina Lee is a sociologist of evaluation, knowledge, and inequality. Her work asks: how do systems of evaluation determine whose contributions are recognized as original, authoritative, and credible? And how does inequality concentrate in the judgments least governed by clear standards? She pursues these questions across three arenas: the sciences; cultural canonization; and the AI systems now absorbing parts of scholarly judgment itself, using computational text analysis, bibliometric analysis, and experiments. Her research appears in the American Sociological Review, Gender & Society (forthcoming), and Poetics, among other venues.":
        "평가, 지식, 불평등을 연구하는 사회학자 이진아입니다. 평가 체계는 누구의 기여를 독창적이고 권위 있으며 신뢰할 만한 것으로 인정할까요? 명확한 기준의 영향을 가장 덜 받는 판단에서 불평등이 왜 집중되는지도 살핍니다. 과학, 문화적 정전화, 그리고 학술적 판단의 일부를 흡수하고 있는 AI 시스템이라는 세 영역을 가로질러 이 질문을 탐구합니다. 전산 텍스트 분석, 서지 분석, 실험을 활용해 여성 연구자와 지위가 낮은 연구자에게 권위와 인정이 덜 부여되며, 그 격차가 기준이 가장 모호한 곳에서 확대된다는 점을 보여줍니다. 연구는 American Sociological Review, Gender & Society (게재 예정), Poetics를 비롯한 여러 학술지에 실렸습니다.",
    "My research asks how evaluation confers authority on some contributions but not others, and why inequality is greatest where judgments are least standardized. For instance, a paper may be widely cited without being treated as original, credible, or field-defining; status inequality accumulates in these harder-to-standardize forms of recognition. I pursue this question across three sites: science, culture, and AI.":
        "평가를 통해 어떤 기여에는 권위가 부여되고 다른 기여에는 그렇지 않은 이유, 그리고 판단이 가장 덜 표준화된 곳에서 왜 불평등이 가장 크게 나타나는지 질문합니다. 예를 들어, 어떤 논문은 널리 인용되면서도 독창적이거나 신뢰할 만하거나 해당 분야를 대표하는 연구로 받아들여지지 않을 수 있습니다. 지위 불평등은 바로 이러한 표준화하기 어려운 인정의 형태 속에서 누적됩니다. 이 질문을 과학, 문화, AI라는 세 영역에서 탐구합니다.",
    "Jina Lee is a sociologist who studies how seemingly objective evaluation systems reproduce social hierarchies in science and in culture. Across these contexts, her work reveals a consistent pattern: practices that appear meritocratic often embed biases that disadvantage women and lower-status actors. Her research has been published in the American Sociological Review, Poetics, Socius, and Journal of Social Entrepreneurship, and is forthcoming in Gender & Society.":
        "저는 사회학자로서 과학과 문화 영역에서 겉보기에는 객관적인 평가 체계가 사회적 위계를 어떻게 재생산하는지 연구합니다. 서로 다른 맥락을 가로질러 살펴보면, 일관된 패턴이 드러납니다. 능력에 기반한 것처럼 보이는 평가나 보상에도 여성과 지위가 낮은 행위자에게 불리하게 작용하는 편향이 내재되어 있습니다. 연구 결과는 American Sociological Review, Poetics, Socius, Journal of Social Entrepreneurship에 게재되었으며, Gender & Society에 게재 예정입니다.",
    "I teach undergraduate courses in sociology of culture, sociology of gender, social statistics, and technology and society. My courses emphasize critical thinking and the application of sociological frameworks to contemporary empirical questions.":
        "문화사회학, 젠더사회학, 사회통계학, 과학기술과 사회 등 다양한 학부 과목을 가르칩니다. 수업에서는 비판적으로 사고하고 사회학적 분석 틀을 오늘날의 경험적 질문에 적용하는 능력을 강조합니다.",
    "My research examines how evaluation systems that appear objective reproduce gender hierarchies across scientific and cultural fields. I ask: whose contributions are recognized as valuable, and whose are discounted? Across these programs, I trace how gender bias operates through everyday evaluation practices and accumulates into durable inequalities.":
        "객관적으로 보이는 평가 체계가 과학과 문화의 여러 영역에서 젠더 위계를 어떻게 재생산하는지 분석합니다. 어떤 기여가 가치 있는 것으로 인정되고, 어떤 기여가 평가절하되는지 질문합니다. 일상적인 평가 관행 속에서 젠더 편향이 작동하고 지속적인 불평등으로 축적되는 과정도 추적합니다.",
    "My methods include computational text analysis, bibliometric analysis, and survey experiments.":
        "전산 텍스트 분석, 서지 분석, 설문 실험을 활용합니다.",
    "PROGRAM: Gender and Scientific Evaluation": "PROGRAM: 젠더와 과학적 평가",
    "Scientific recognition is rarely neutral. My work in this area investigates how gender shapes which knowledge claims are treated as authoritative, how novelty is attributed and rewarded, and how uncertainty is managed differently depending on who makes a claim.":
        "과학적 인정은 중립적으로 이루어지지 않습니다. 젠더가 어떤 지식 주장을 권위 있는 것으로 만들고, 새로움을 누구에게 귀속·보상하며, 과학적 주장을 누가 제시하느냐에 따라 불확실성을 어떻게 다르게 관리하는지 분석합니다.",
    "How do gendered dynamics shape which scientific contributions get recognized, cited, and treated as authoritative?":
        "젠더화된 평가 관행은 어떤 과학적 기여가 인정되고, 인용되며, 권위 있는 것으로 받아들여지는지를 어떻게 형성하는가?",
    "How is novelty attributed and rewarded differently across gender lines?":
        "과학적 새로움은 젠더에 따라 어떻게 다르게 귀속되고 보상되는가?",
    "How is uncertainty managed differently depending on who makes a scientific claim?":
        "과학적 주장을 누가 제시하느냐에 따라 불확실성은 어떻게 다르게 관리되는가?",
    "PROGRAM: Gender, Culture, and Recognition": "PROGRAM: 젠더, 문화, 그리고 인정",
    "Recognition unfolds across cultural fields, markets, and public arenas. My work in this area examines how gender and social hierarchies shape who is recognized, remembered, funded, protected, or able to remain publicly present.":
        "인정은 문화적 장, 시장, 공적 영역을 가로질러 형성됩니다. 젠더와 사회적 위계가 누가 인정받고 기억되며, 자금을 지원받고 보호받거나, 공적 공간에 계속 남을 수 있는지를 어떻게 형성하는지 분석합니다.",
    "How do cultural institutions determine whose work is recognized and remembered?":
        "문화 제도는 누구의 작업이 인정되고 기억될지를 어떻게 결정하는가?",
    "How do gender stereotypes shape access to recognition and resources?":
        "젠더 고정관념은 인정과 자원에 대한 접근을 어떻게 형성하는가?",
    "How do social categories influence judgments of who deserves support or protection?":
        "사회적 범주는 누가 지원이나 보호를 받을 자격이 있는지에 대한 판단에 어떤 영향을 미치는가?",
    "At which stages of evaluation do inequalities emerge and accumulate?":
        "불평등은 평가의 어느 단계에서 발생하고 축적되는가?",
    "PROGRAM: Science and Academia": "PROGRAM: 과학과 학계",
    "Scientific knowledge is shaped not only by what researchers study but by how scientific communities are organized. This line of work examines how structural features of academic fields, such as how audiences are structured, publication norms, and data practices, influence the production, reception, and epistemic character of scientific knowledge.":
        "과학적 지식은 연구자들이 무엇을 연구하는지뿐 아니라 과학 공동체가 어떻게 조직되는지에 의해 형성됩니다. 학문 분야의 청중 구조, 출판 규범, 데이터 관행과 같은 구조적 특성이 과학 지식의 생산과 수용, 인식론적 성격에 어떤 영향을 미치는지 분석합니다.",
    "How does audience structure shape the impact of domain-spanning innovation?":
        "청중의 구조는 여러 영역을 가로지르는 혁신의 영향력을 어떻게 형성하는가?",
    "How do editorial and data sharing requirements shape the epistemic character of scientific articles?":
        "편집 및 데이터 공유 요건은 과학 논문의 인식론적 성격을 어떻게 형성하는가?",
    "PROGRAM: AI and Research Practice": "PROGRAM: AI와 연구 실천",
    "A newer line of my work asks how AI tools are changing research practice itself, not just its speed or volume.":
        "최근에는 AI 도구가 연구의 속도나 규모를 넘어 연구 관행 자체를 어떻게 변화시키는지 살펴봅니다.",
    "How can AI-assisted research systems make interpretive decisions visible and auditable?":
        "AI 지원 연구 시스템은 해석적 결정을 어떻게 가시화하고 검증·감사할 수 있게 만들 수 있는가?",
    "What survives, and what is lost, when human judgment is relocated into external infrastructure a researcher runs alone?":
        "연구자가 혼자 실행하는 외부 인프라로 인간의 판단이 이전될 때, 무엇이 남고 무엇이 사라지는가?",
    "I believe in the transformative power of integrative and experiential learning to make sociological concepts tangible and relevant to students' lives. I design interactive learning environments where theory comes alive through collaborative investigation and creative application. Below are select examples from three of my undergraduate courses.":
        "통합적·경험적 학습은 사회학적 개념을 학생들의 삶과 연결해 구체적으로 이해하도록 돕습니다. 저는 협력적 탐구와 창의적 적용을 통해 이론을 실제 문제에 활용하는 참여형 학습 환경을 설계합니다. 아래에는 제가 학부 세 과목에서 활용한 대표 활동을 소개합니다.",
    "A survey of fundamental statistical concepts and their application in social research, with emphasis on both conceptual understanding and practical analysis skills.":
        "사회 연구에 필요한 기초 통계 개념과 활용법을 다루며, 개념적 이해와 실제 자료 분석 능력을 함께 기릅니다.",
    "Key frameworks in the sociology of culture, including cultural industries, taste, status, and the boundary between popular and high culture. Students analyze everyday cultural objects through a sociological lens.":
        "문화산업, 취향, 지위, 대중문화와 고급문화의 경계 등 문화사회학의 주요 분석 틀을 다룹니다. 학생들은 일상에서 접하는 문화적 대상을 사회학적 관점으로 분석합니다.",
    "How gender shapes experience across education, work, family, and relationships, and how it intersects with other dimensions of social inequality.":
        "젠더가 교육, 일, 가족, 관계에서의 경험을 어떻게 형성하는지, 그리고 사회 불평등의 다른 차원과 어떻게 교차하는지를 다룹니다.",
    "How science and technology shape social life, and how social forces shape science and technology in return. Topics include recognition, inequality, AI governance, and algorithmic systems.":
        "과학과 기술이 사회적 삶을 어떻게 형성하는지, 반대로 사회적 힘은 과학과 기술을 어떻게 형성하는지를 살펴봅니다. 인정, 불평등, AI 거버넌스, 알고리즘 시스템 등을 주요 주제로 다룹니다.",
    "Students map their campus environment to identify how institutional resources are distributed along class-coded lines, drawing on Anthony Jack's research on first-generation students.":
        "학생들은 캠퍼스를 직접 살펴보며 계층적 신호에 따라 제도적 자원이 어떻게 배분되는지 찾아내고, Anthony Jack의 가족 중 처음으로 대학에 진학한 학생들에 관한 연구를 바탕으로 그 패턴을 분석합니다.",
    "Teams represent stigmatized subcultures and use theoretical tools to argue for their group's legitimacy, navigating between analytical rigor and rhetorical performance.":
        "각 팀은 낙인찍힌 하위문화 집단을 맡아 이론적 도구를 활용해 해당 집단의 정당성을 주장합니다. 이 과정에서 분석적 엄밀성과 설득력 있는 표현 사이의 균형을 익힙니다.",
    "Students translate sociological concepts into visual form and analyze what makes a concept-driven meme culturally legible and shareable.":
        "학생들은 사회학적 개념을 시각적으로 표현하고, 개념 중심의 밈이 문화적으로 쉽게 읽히고 공유될 수 있는 조건을 분석합니다.",
    "Students construct evidence-based counter-memes to challenge essentialist claims, practicing critical engagement with popular discourse about gender.":
        "학생들은 본질주의적 주장을 반박하는 근거 기반의 밈을 만들며, 젠더에 관한 대중 담론을 비판적으로 검토하는 법을 연습합니다.",
    "Students curate visual representations of gendered professional norms and present them analytically, surfacing the social regulation of bodies in professional contexts.":
        "학생들은 직업 세계의 젠더 규범을 보여주는 시각 자료를 선별·구성해 분석적으로 발표하고, 직업적 맥락에서 신체가 사회적으로 규율되는 방식을 드러냅니다.",
    "Groups write fictional letters from 1950s time travelers confused by modern work-family arrangements; peers respond with sociologically informed analysis connecting personal experience to institutional change.":
        "각 조는 오늘날의 일과 가족생활 방식을 낯설어하는 1950년대 시간 여행자의 가상 편지를 작성합니다. 동료 학생들은 개인적 경험과 제도적 변화를 연결하는 사회학적 분석으로 답합니다.",
    "Students use structured gameplay to examine how scientific credit accumulates unequally, mapping advantage and disadvantage onto the social conditions of knowledge production.":
        "학생들은 구조화된 게임을 통해 과학적 공로가 불평등하게 축적되는 방식을 살펴보고, 지식 생산의 사회적 조건이 누구에게 유리하거나 불리하게 작용하는지를 파악합니다.",
    "Students document their everyday encounters with algorithmic systems, then analyze them as a class to surface patterns in how AI shapes choice and perception.":
        "학생들은 일상에서 알고리즘 시스템을 접한 경험을 기록하고, 이를 수업에서 함께 분석해 AI가 선택과 인식을 형성하는 공통된 양상을 찾아냅니다.",
    "Groups construct physical representations of differential technology access, making structural inequality in digital infrastructure materially visible.":
        "학생들은 기술 접근성의 차이를 물리적 모형으로 표현해 디지털 인프라에 내재한 구조적 불평등을 눈에 보이게 만듭니다.",
}

KOREAN_FONT: Final = "Apple SD Gothic Neo"


class MissingKoreanTranslationError(RuntimeError):
    """Raised when source prose has no reviewed Korean translation."""


PRESERVED_TITLES = {
    "Claiming Novelty, Claiming Authority: Gender Gaps in Scientific Impact Across Disciplines",
    "The Theory Penalty: Gender Bias in Recognition of Scientific Novelty",
    "Stratified Fact-Making: How Gender and Novelty Claims Stratify the Stabilization of Scientific Facts",
    "What Types of Novelty Are Most Disruptive?",
    "Gendered Pathways to Perpetual Fame: The Selection of Elite Novelists into the Korean Literary Canon",
    "Who Deserves Protection? How Naming Potential Beneficiaries Influences the COVID-19 Vaccine Intentions",
    "The Persistent Influence of Gender Stereotypes in Social Entrepreneurial Financing",
    "Survivorship in Public: Differential Durability and the Conflictual Face of Korean Digital Feminism",
    "Divide and Conquer? How Partitioned Audiences Shape the Impact of Domain-Spanning Innovation",
    "Humble Reflections on the Intellectual Process of Developing a Text-based Measure of Humility in Inquiry",
    "Conceptual Divergence Analysis: Mapping a Researcher's Conceptual Vocabulary Against the Literatures They Address",
}


def normalize_korean_translation(text: str) -> str:
    """Keep the author's established Korean name in translated site prose."""
    return text.replace("Jina Lee", "이진아")


def should_preserve(text: str) -> bool:
    return text in PRESERVED_TITLES or text.startswith((
        "HOME ",
        "RESEARCH INTRO",
        "KEY QUESTIONS",
        "PUBLICATIONS",
        "TEACHING PHILOSOPHY",
        "COURSE: ",
        "ACTIVITIES: ",
        "ITEM: ",
    )) or text in {
        "HOME ABOUT",
        "HOME TEACHING",
        "PUBLICATIONS",
        "TEACHING PHILOSOPHY",
    }


def replace_paragraph_text(paragraph: Paragraph, replacement: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(replacement)
    else:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ""

    for run in paragraph.runs:
        run.font.name = KOREAN_FONT
        run._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), KOREAN_FONT
        )


def create_korean_document(
    input_path: Path,
    output_path: Path,
    auto_translations_path: Path,
    model: str,
) -> None:
    document = Document(input_path)
    auto_translations = load_auto_translations(auto_translations_path)
    new_paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if (
            text
            and text not in TRANSLATIONS
            and text not in auto_translations
            and not should_preserve(text)
            and text not in new_paragraphs
        ):
            new_paragraphs.append(text)

    if new_paragraphs:
        generated = {
            source: normalize_korean_translation(translation)
            for source, translation in translate_prose(new_paragraphs, model).items()
        }
        auto_translations.update(generated)
        save_auto_translations(auto_translations_path, auto_translations)
        print(
            f"Auto-translated {len(generated)} new paragraphs with {model}; "
            f"review them in {auto_translations_path}."
        )

    translated = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        replacement = TRANSLATIONS.get(text) or auto_translations.get(text)
        if replacement is not None:
            replace_paragraph_text(paragraph, normalize_korean_translation(replacement))
            translated += 1
        elif should_preserve(text):
            continue
        else:
            raise MissingKoreanTranslationError(
                f"No Korean translation generated for: {text}"
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Created {output_path} with {translated} translated paragraphs.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create Korean site-content.docx")
    parser.add_argument("--input", default="site-content.docx")
    parser.add_argument("--output", default="site-content-ko.docx")
    parser.add_argument(
        "--auto-translations",
        default=str(DEFAULT_AUTO_TRANSLATIONS_PATH),
    )
    parser.add_argument(
        "--model",
        default=os.environ.get("OPENAI_TRANSLATION_MODEL", DEFAULT_TRANSLATION_MODEL),
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    create_korean_document(
        Path(args.input),
        Path(args.output),
        Path(args.auto_translations),
        args.model,
    )


if __name__ == "__main__":
    main()
